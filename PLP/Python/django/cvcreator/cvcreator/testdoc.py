from django.http import HttpResponseRedirect, HttpResponse
from docx import *
from docx.shared import Inches
#import pdfkit
import io


def testDocument():
    document = Document()
    docx_title = "TEST_DOCUMENT.docx"
    pdf_title = "TEST_DOCUMENT. pdf"

    # __CV doc__ #
    p = document.add_paragraph()
    p.add_run("Test text1.")
    p.add_run("Test text2.")
    document.add_page_break()

    # Prepare the document for download #
    f = io.BytesIO()
    print(f)
    document.save(f)
    length = f.tell()
    f.seek(0)

    '''
    #____ pdf format also____#
    options = {
        'page-size':'Letter',
        # 'format':'pdf',
        'encoding':'utf-8',
    }

    pdf = pdfkit.from_string(f.getvalue(),False, options)
    response = HttpResponse(pdf, content_type="application/pdf")
    response['Content-Disposition'] = "attachment; filename=" + pdf_title
    response["Content-Length"] = length
    '''

    # ____ word docx format ____#
    response = HttpResponse(f.getvalue(
    ), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response['Content-Disposition'] = "attachment; filename=" + docx_title
    response["Content-Length"] = length

    return response


testDocument()