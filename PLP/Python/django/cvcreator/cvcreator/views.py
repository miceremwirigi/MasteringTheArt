import io
import collections
import pdfkit as pdf
import random
from django.shortcuts import render, redirect

# Create your views here.

from .models import PersonalInfo
from .forms import PersonalInfoForm
from django.http import HttpResponseRedirect, HttpResponse
from django.urls import reverse
from docx import *
from docx.shared import Inches


def cvcreator(request):
    return render(request, "cvcreator.html")

# view to show members in the database


def listCustomers(request):
    members = PersonalInfo.objects.all()

    context = {
        'members': members,
    }
    return render(request, 'listcustomers.html', context)


# Form to add customer info to database
def createnewinfo(request):
    form = PersonalInfoForm()
    if request.method == "POST":
        form = PersonalInfoForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return HttpResponseRedirect(reverse("cvcreator"))

    context = {
        "form": form,
    }
    return render(request, "createnewinfo.html", context)


def deleteCustomer(request, pk):
    customer = PersonalInfo.objects.get(idNumber=pk)
    customer.delete()
    return HttpResponseRedirect(reverse("cvcreator"))

# Show details of all information of specific customers, e.g., name, ID, phone etc.


def customerInfo(request, pk):
    infos = PersonalInfo.objects.filter(idNumber=pk).values()[0]

    context = {
        "infos": infos,
    }
    print(infos)
    print(type(infos))
    for info, value in infos.items():
        print(info)
        print(value)
    return render(request, "customerinfo.html", context)


# generate and download specific CV document
def getDocument(request, pk):
    infos = PersonalInfo.objects.filter(idNumber=pk).values()[0]
    document = Document()
    print(infos)
    docx_title = infos.get('name') + "_CV.docx"
    pdf_title = "TEST_DOCUMENT. pdf"

    # __CV doc__ #
    h = document.add_heading()
    h.add_run(u"Personal Info: \n", 0)
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 2
    # After converting to ordered dict above, we can check for last item in function below

# ______***All this def lastItem() hustle to remove one empty line from word document ***______#
# _____________________________________***ha ha***____________________________________________#

    def lastItem(myModel, key):
        infos = myModel.objects.filter(idNumber=pk).values()[0]
        # convert to ordered dict to ensure no  blank line is left after te last item
        infos = collections.OrderedDict(infos)
        print(infos.items())
        # fieldName = list(infos.items())[-1][0]
        # print(fieldName)  # get a list of key value pairs of this record/instance

        # Getting the name of the field last field
        pi = myModel()  # create a model instance
        all_fields = pi._meta.fields
        lastField = all_fields[-1].name
        print(lastField)  # Get the name pf the last fieldin the model
        if key == lastField:
            return True
        else:
            return False
    # ___________________________________***hu hu hu!***________________________________________#

    for info, value in infos.items():
        if info == "idNumber":
            info = "ID Number"
        elif info == "maritalStatus":
            info = "Marital Status"
        elif info == "dob":
            info = "Born"
        
        if lastItem(PersonalInfo, info):
            p.add_run(info.title()). bold = True
            p.add_run(":      " + str(value) + ".")
        else:
            p.add_run(info.title()). bold = True
            p.add_run(":      " + str(value) + ". \n")

    # ______________________________*** Career Objectives ***___________________________________#
    skill_list = ["To advance my understanding, expertise, and abilities by obtaining a demanding position in a recognized company.",
                  "Obtain a fantastic professional position that will allow me to make the most of my education and experience while significantly contributing to the company's development.",
                  "Looking for an entry-level job to start my career in a very professional setting.",
                  "To secure employment with a reputable company, where I can utilize my skills and business studies background to the maximum.",
                  "A diligent and well-organized person seeking a position of responsibility to build experience",
                  "To make use of my interpersonal skills to achieve goals of a company that focuses on customer satisfaction and customer experience.",
                  "Strive to meet the highest standards expected of me by my employer, and be self-driven and task-oriented in completing the obligations placed upon me.",
                  "Promote personal and professional progress for myself and the organization.",
                  "I'm looking for a competitive, demanding environment where I can work for a company and build a career.",
                  ]
    
    
    h = document.add_heading()
    h.add_run(u"Career Objective \n", 0)
    p.paragraph_format.line_spacing = 2

    count = 0
    while count < 2:
        p = document.add_paragraph()
        p.style = 'List Bullet'
        random.shuffle(skill_list)
        skill = random.choice(skill_list)
        p.add_run(skill)
        count += 1
    # ______________________________*** Career Objectives End***___________________________________#

    h = document.add_heading()
    h.add_run(u" Skills", 0)

    # ______________________________*** Skills ***___________________________________#
    skill_list = ["Computer proficiency",
                  "Leadership experience",
                  "Communication skills",
                  "Organizational know-how",
                  "People skills",
                  "Collaboration talent",
                  "Problem-solving abilities",
                  "Typing",
                  "Bookkeeping",
                  "Customer service",
                  "Clerical Expertise",
                  "Document management",
                  "Employee relations",
                  "Filling",
                  "Event coordination",
                  "Calendar docketing",
                  "Appointment setting",
                  "Microsoft office",
                  "Office administration",
                  "Receptionist",
                  "Written communication",
                  "Content management systems(CMS)",
                  "Foreign languages",
                  "Operating certain equipment or machinery",
                  "Patient preparation",
                  "Point-of-sale(POS) system",
                  "Project management",
                  "Software Development",
                  "Scheduling",
                  "Bookkeeping",
                  "Content management systems(CMS)",
                  "Foreign languages",
                  "Operating certain equipment or machinery",
                  "Patient preparation",
                  "Point-of-sale(POS) system",
                  "Project management",
                  "Software Engineering",
                  "Scheduling",
                  "Taking vital signs",
                  "Communication",
                  "Customer service",
                  "Decision-making",
                  "Driving",
                  "Integrity",
                  "Leadership",
                  "Organization",
                  "Problem-solving",
                  "Teamwork",
                  "Time management"]

    count = 0
    while count < 7:
        p = document.add_paragraph()
        p.style = 'List Bullet'
        random.shuffle(skill_list)
        skill = random.choice(skill_list)
        p.add_run(skill)
        count += 1

    document.add_page_break()
    # ______________________________*** Skills End***___________________________________#
    # End of page 1

    # ______________________________*** Personal Profile ***___________________________________#
    profile_list = ["I'm an enthusiastic, self-motivated, reliable, responsible and hard working \person.",
                    "I'm a mature team worker and adaptable to all challenging situations.",
                    "I am able to work well both in a team environment as well as using own initiative.",
                    "I'm able to work well under pressure and adhere to strict deadlines.",
                    "I am a punctual and motivated individual who is able to work in a busy environment and produce high standards of work.",
                    "I work well in a team, can take direction at all levels, and develop strong working connections with all of my coworkers.",
                    "I'm adaptable, trustworthy, and very good at keeping my word.",
                    "I am a diligent and trustworthy person.",
                    "I keep good track of time and am always eager to pick up new abilities.",
                    "I have an excellent sense of humor and am amiable, helpful, and courteous.",
                    "I am a dedicated and sincere person. I am dependable with my time and open to picking up new abilities.",
                    "I am affable, kind, and courteous, and I also have a fantastic sense of humor.",
                    "My approach to work and accomplishing goals is energetic and dynamic.",
                    "I have a strong sense of purpose. ", "I locate opportunities and create them.",
                    "I'm motivated to see things through to the end and have a clear, logical mind. ",
                    "I also have a practical approach to problem-solving.",
                    "I have a genuine interest in business management and the development of organizations, I am willing to learn, and I like conquering obstacles.",
                    "I can engage with a variety of clientele because I have outstanding interpersonal skills in addition to strong technical abilities.",
                    "I'm an energetic, ambitious person who has developed a mature and responsible approach to any task that I undertake, or situation that I am presented with.",
                    "I'm an experienced, upbeat, and diligent person that always aims to meet the greatest quality attainable in each given endeavor.",]

    h = document.add_heading()
    h.add_run(u"Personal Profile\n", 0)
    count = 0
    p = document.add_paragraph()
    while count < 5:
        random.shuffle(profile_list)
        skill = random.choice(profile_list)
        if count == 0:
            p.add_run(skill)
        else:
            p.add_run(" " + skill)
        count += 1
    # ______________________________***Personal Profile End***___________________________________#

    # Prepare the document for download #
    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)

    '''
    #____ pdf format also____#
    options = {
        'page-size':'Letter',
        # 'format':'pdf',
        'encoding':'utf-8',
    }

    pdf = pdfkit.from_string(f.getvalue(),False, options)
    response = HttpResponse(pdf, content_type="application/pdf")
    response['Content-Disposition'] = "attachment; filename=" + pdf_title
    response["Content-Length"] = length
    '''

    # ____ word docx format ____#
    response = HttpResponse(f.getvalue(
    ), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response['Content-Disposition'] = "attachment; filename=" + docx_title
    response["Content-Length"] = length

    return response


# # to test how to work wuth docx
# def testDocument(request):
#     document = Document()
#     docx_title = "TEST_DOCUMENT.docx"
#     pdf_title = "TEST_DOCUMENT. pdf"

#     #__CV doc__ #
#     p = document.add_paragraph()
#     p.add_run("Test text1.")
#     p.add_run("Test text2.")
#     document.add_page_break()

#     # Prepare the document for download #
#     f = io.BytesIO()
#     print(f)
#     document.save(f)
#     length = f.tell()
#     f.seek(0)

#     '''
#     #____ pdf format also____#
#     options = {
#         'page-size':'Letter',
#         # 'format':'pdf',
#         'encoding':'utf-8',
#     }

#     pdf = pdfkit.from_string(f.getvalue(),False, options)
#     response = HttpResponse(pdf, content_type="application/pdf")
#     response['Content-Disposition'] = "attachment; filename=" + pdf_title
#     response["Content-Length"] = length
#     '''

#     # ____ word docx format ____#
#     response = HttpResponse(f.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
#     response['Content-Disposition'] = "attachment; filename=" + docx_title
#     response["Content-Length"] = length


#     return response

# Show all clients whose data is available in the database
